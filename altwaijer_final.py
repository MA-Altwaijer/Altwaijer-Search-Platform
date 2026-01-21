import streamlit as st
import pandas as pd
import google.generativeai as genai
from docx import Document
from io import BytesIO
from pypdf import PdfReader

# 1. الربط الآمن بخزنة الأسرار (هذا ما ينقص المنصة الآن)
try:
    api_key = st.secrets["GEMINI_API_KEY"]
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-pro')
except:
    st.error("⚠️ يرجى التأكد من ضبط المفتاح السري في إعدادات Secrets")

st.set_page_config(page_title="Altwaijer Hub", layout="wide")
st.markdown("<h1 style='text-align:center; color: #1E3A8A;'>🏛️ منصة M.A. Altwaijer للتميز والابتكار</h1>", unsafe_allow_html=True)

# دالة استخراج النص من الـ PDF المرفوع
def extract_text(files):
    all_text = ""
    for f in files:
        reader = PdfReader(f)
        for page in reader.pages[:3]: # قراءة أول 3 صفحات للسرعة والدقة
            all_text += page.extract_text()
    return all_text

# واجهة المستخدم
st.sidebar.header("🎯 مسار بناء البحث")
step = st.sidebar.radio("المراحل المنهجية:", ["1. تحديد العنوان", "2. صياغة الإطار النظري", "3. تحميل المسودة"])

files = st.file_uploader("📂 ارفعي المراجع (PDF):", type="pdf", accept_multiple_files=True)

if files:
    with st.spinner("⏳ جاري تحليل المحتوى العلمي..."):
        context = extract_text(files)
        
        if step == "1. تحديد العنوان":
            st.subheader("💡 مقترحات عناوين بحثية ذكية:")
            prompt = f"بناءً على هذا النص: {context[:4000]}، اقترح 3 عناوين بحثية مبتكرة ورصينة."
            response = model.generate_content(prompt)
            st.write(response.text)

        elif step == "2. صياغة الإطار النظري":
            st.subheader("📝 صياغة أكاديمية مقترحة (APA):")
            prompt = f"بناءً على الدراسات المرفقة: {context[:4000]}، اكتب فقرة إطار نظري تربط بين النتائج مع التوثيق."
            response = model.generate_content(prompt)
            st.session_state['theory'] = response.text
            st.write(response.text)

        elif step == "3. تحميل المسودة":
            if 'theory' in st.session_state:
                doc = Document()
                doc.add_heading("مسودة الإطار النظري - د. مبروكة التويجر", 0)
                doc.add_paragraph(st.session_state['theory'])
                buffer = BytesIO()
                doc.save(buffer)
                st.download_button("📥 تحميل ملف Word المنسق", buffer.getvalue(), "Altwaijer_Draft.docx")
            else:
                st.warning("يرجى الانتقال للمرحلة الثانية أولاً لتوليد النص.")

st.markdown("---")
st.caption("إشراف وتطوير: د. مبروكة التويجر - 2026")
